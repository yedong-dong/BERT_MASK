
import os
import re
import string
from fuzzywuzzy import fuzz
from transformers import BertTokenizer, BertForMaskedLM, BertForPreTraining
from keyword_extractor import chinese_tokenizer ,extract_keywords
import torch
import re
import random
import os
import warnings
from document_reader import read_document_by_pages
from docx import Document as DocxDocument
from docx.shared import RGBColor
class DocumentProcessor:
    def __init__(self, file_path):
        self.file_path = file_path

    def process_document(self):
        try:
            pages_content = read_document_by_pages(self.file_path)

        except Exception as e:
            print(f"读取文件时出错: {e}")
            return

        # 打印每一页的内容
        ContentList=[]
        for i, page in enumerate(pages_content):
            print(f"Page {i + 1}:")
            print(page)
            ContentList.append(page)
        return ContentList

# def load_whiteList():
#     with open("./whiteList",'a', encoding="utf-8") as f:
#         for line in f:
#             whiteList.append(line.strip())
#     return whiteList
def load_model():
    global tokenizer, model
    warnings.filterwarnings("ignore")
    # 加载已保存的 BERT 分词器和模型
    tokenizer = BertTokenizer.from_pretrained('03_trained_bert')
    model = BertForMaskedLM.from_pretrained('03_trained_bert')
    correct_predictions_file = 0
    total_sentences_file = 0
def readContent(read_filename,file_type,input_type):
    # 提示用户选择文件类型
    #  file_type = input("请选择文件类型 (docx/doc/pdf): ").strip().lower()

    # 提示用户选择输入方式
    # input_type = input("请选择输入方式 (name/path): ").strip().lower()

    # 根据用户选择的输入方式获取文件路径
    if input_type == "name":
        # 提示用户输入文件名
        file_name = input("请输入同目录下的文件名(无后缀): ").strip()

        # 获取当前工作目录
        current_directory = os.getcwd()
        # 构建文件的绝对路径
        file_path = os.path.join(current_directory, file_name + "." + file_type)
    elif input_type == "path":
        # 提示用户输入文件的完整路径
        # file_path = input("请输入文件的完整路径(无双引号): ").strip()
        file_path=read_filename.strip().lower()

    else:
        print("不支持的输入方式，请选择 name 或 path。")
        return

    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"文件 {file_path} 不存在。")
        return

    # 检查文件类型和文件路径是否匹配
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension != "." + file_type:
        print(f"文件类型和文件路径不匹配。选择的文件类型是 {file_type}，但文件路径的扩展名是 {file_extension}。")
        return

    # 创建 DocumentProcessor 实例并处理文档
    processor = DocumentProcessor(file_path)
    return processor.process_document()

# 预测
def predict(Contents,whitelist):
    # 去掉换行符和特殊字符
    chinese_punctuation = '，。！？“”；：（）《》【】‘’ [UNK]'
    append_error=[]
    cleaned_contents = []
    iter = 0
    correct_predictions_file=0
    total_sentences_file=0
    result=[]
    result_most=[]

    for content in Contents:
        cleaned_content = content.replace('\n', '').replace('\u3000', '').strip()  # 清洗换行符和全角空格
        cleaned_content = cleaned_content.replace(' ', '')  # 删除空格
        cleaned_contents.append(cleaned_content)  # 将清洗后的内容添加到 cleaned_contents 列表中

    print(cleaned_contents)
    # 分段判断
    # para_list = cleaned_contents[0].split('\r')
    # print(para_list)

    for para_Content in cleaned_contents:
        # 第n段落的句子
        sentences = re.split(r'[。；]', para_Content)
        iter += 1
        # 进行遍历第一页的句子

        quantifiers = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']

        for paraInner in range(len(sentences)):
            if len(sentences[paraInner]) > 6:
                sentence_first = sentences[paraInner][:6] + "..."
            else:
                sentence_first = paraInner  # 如果长度小于或等于6，直接使用原字符串
            sentence=sentences[paraInner]
            paraInner=paraInner+1
            sentence = sentence.strip()
            # sentence=sentence.replace(' ','')
            if sentence:  # 确保句子不为空
                chars_length = list(sentence)
                for j in range(len(chars_length)):
                    chars = list(sentence)
                    if (chars[j] in string.punctuation
                            or chars[j] in string.ascii_letters\
                            or chars[j] in chinese_punctuation\
                            or chars[j].isdigit())\
                            or chars[j] in quantifiers:continue

                    if random.random() <0.04:
                        chars[j] = '错'
                        append_error.append(f"{iter}-{paraInner}-{j + 1}")
                        print(f"{iter}-{paraInner}-{j + 1}个词,改为错")

                    mask_position = j
                    original_word = chars[mask_position]
                    chars[mask_position] = tokenizer.mask_token  # 替换为掩码
                    masked_sentence = ''.join(chars)
                    inputs = tokenizer(masked_sentence, return_tensors='pt')

                    with torch.no_grad():
                        outputs = model(**inputs)
                        errors = []  # 用于存储未能正确预测的单词
                        logits = outputs.logits
                        mask_token_index = torch.where(inputs['input_ids'] == tokenizer.mask_token_id)[1]
                        mask_logits = logits[0, mask_token_index, :]
                        top_k = 5
                        top_k_scores, top_k_indices = torch.topk(mask_logits, top_k)
                        probabilities = torch.nn.functional.softmax(mask_logits, dim=-1)

                        print(f"句子：{masked_sentence}")
                        print(f"被掩码的原始词：{original_word}")
                        print("可能的填空词及其得分：")
                        predicted_tokens=[]
                        predicted_probability_list=[]
                        for score, index in zip(top_k_scores[0], top_k_indices[0]):
                            predicted_token = tokenizer.decode([index])
                            predicted_probability = probabilities[0, index].item()
                            predicted_tokens.append(predicted_token)
                            predicted_probability_list.append(round(predicted_probability, 4))
                            print(f"词：{predicted_token}, 得分：{score.item()}, 概率：{predicted_probability:.4f}")


                        # 解码并检查准确性
                        predicted_words = [tokenizer.decode([index]) for index in top_k_indices[0]]
                        is_correct = original_word in predicted_words[:5]

                        total_sentences_file += 1
                        if is_correct:
                            correct_predictions_file += 1
                        # else:
                        #     if mask_position != 0 and mask_position != len(sentence) - 1:
                        #         original_sentence = sentence[mask_position - 1:mask_position + 2]
                        #     elif mask_position == 0:  # 如果掩码字是句子的第一个字
                        #         original_sentence = sentence[mask_position:mask_position + 3]
                        #     elif mask_position == len(sentence) - 1:  # 如果掩码字是句子的最后一个字
                        #         original_sentence = sentence[mask_position - 2:mask_position + 1]
                        #
                        #     found_similar_word = False  # 标记是否找到相似白名单词
                        #     for word in whitelist:
                        #         # 使用模糊匹配检查原始句子与白名单词汇的相似度
                        #         if fuzz.partial_ratio(original_sentence, word) > 75:
                        #             found_similar_word = True  # 找到相似的白名单词
                        #             break  # 找到相似的白名单词则跳出循环
                        else:
                            # if not found_similar_word:
                             symbol_count = sum(1 for token in predicted_token if token in string.punctuation)
                             if symbol_count < 3:
                                  # 如果没有找到相似的白名单词，则记录错误
                                  print(
                                      f"第{iter}段的第{paraInner}个句子的第{j + 1}词'{original_word}'可能存在错误，预测结果推荐从高到低为：{predicted_tokens}")
                                  result.append(
                                      f"{iter}-{paraInner}-{j + 1}-{sentence_first}-{original_word}:{predicted_tokens}:{predicted_probability_list}")
                             if predicted_probability_list[0] +predicted_probability_list[1]+predicted_probability_list[2]+predicted_probability_list[3]+predicted_probability_list[4] >0.95:
                                 if symbol_count < 3:
                                      result_most.append(f"{iter}-{paraInner}-{j + 1}-{sentence_first}-{original_word}:{predicted_tokens}:{predicted_probability_list}")

                        print()

    print(f"正确率: {correct_predictions_file / total_sentences_file}")
    return result,result_most,append_error
# 输出信息


def write2Info(result,filename="./myOutput.txt"):
    with open(filename, "w", encoding="utf-8") as f:
       for line in result:
           f.write(line + "\n")

# def add2whiteList(lists, file_path="./whiteList"):
#     # 使用 with 语句打开文件，确保文件操作后自动关闭
#     """
#         list的是添加的白名单
#     """
#     with open(file_path, 'a', encoding="utf-8") as file:  # 以追加模式打开文件
#         for item in lists:
#             file.write(str(item) + "\n")  # 确保将 item 转换为字符串
#     print('写入成功')

def load_whitelist(filename):
    whiteList = []
    if os.path.exists(filename):
        # 读取白名单文件
        with open(filename, 'r',encoding='utf-8') as f:  # 确保以只读模式打开
            for line in f:
                whiteList.append(line.strip())
    else:
        print("文件名不符合规定或为None,白名单返回空值")
    return whiteList


def write_errors_to_docx(pages_content, all_errors, file_path):
    doc = DocxDocument()
    for page in pages_content:
        paragraph = doc.add_paragraph()
        for i, char in enumerate(page):
            if (i, char) in all_errors:
                run = paragraph.add_run(char)
                run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                run = paragraph.add_run(char)
    output_file_path = file_path.replace(".doc", "_errors.doc").replace(".docx", "_errors.docx").replace(".pdf", "_errors.docx")
    doc.save(output_file_path)
    print(f"预测错误的字已标红并保存到 {output_file_path}")



def openbert1(filename,whitelist_file,write_to="./myOutput.txt",write_most_to="./myOutputmost.txt"):
    intersection_rate_sum=0
    correct_rate_sum=0

    # 指定要遍历的目录
    directory = r"D:\python\python-Project1\项目\test-trainer\Utils\mydocx"
    # 遍历目录
    fileiter=0
    for filename in os.listdir(directory):
        # 检查是否是文件
        full_path = os.path.join(directory, filename)  # 组合成完整路径
        if os.path.isfile(full_path):
            fileiter=fileiter+1

            whitelist=load_whitelist(whitelist_file)
            load_model()# 加载模型
            Content=readContent(full_path,full_path.split(".")[-1],'path')# 读取文件
            predicts,predicts_most,append_error=predict(Content,whitelist)# 预测的结果（错误）
            # write_errors_to_docx(Content, error, read_filename) #写入word

            write2Info(predicts,write_to)# 写入
            write2Info(predicts_most,write_most_to) # 写入
            write2Info(append_error,"error.txt")
            intersection_rate,correct_rate=compare2ErrorAndOutput(predicts,predicts_most,append_error)

            intersection_rate_sum =intersection_rate_sum+intersection_rate
            correct_rate_sum = correct_rate_sum+correct_rate

    average_intersection_rate=intersection_rate_sum/fileiter
    average_correct_rate=correct_rate_sum/fileiter

    print(f"平均交集率: {average_intersection_rate}")
    print(f"平均正确率: {average_correct_rate}")
def compare2ErrorAndOutput(predicts,predicts_most,append_error):
    common_predict=0
    common_predict_most=0
    print(append_error)
    for predict1 in predicts:
        prefix1='-'.join(predict1.split('-')[:3])
        if any(prefix1 in error for error in append_error):
            common_predict+=1
    # for predict2 in predicts_most:
    #     prefix2='-'.join(predict2.split('-')[:3])
    #     if any(prefix2 in error for error in append_error):
    #         common_predict_most += 1

    # 假设 common_predict 是预测错误的字的数量
    # append_error 是错误字的数量
    # predicts 是预测列表的数量

    # 计算寻错率
    intersection_rate = 0
    correct_rate = 0

    if len(append_error) != 0:
        intersection_rate = (common_predict - 1) / (len(append_error)+1) if common_predict > 1 else 0  # 避免负值
        print(f"predict预测错误的字与错误的字的交集率：{intersection_rate}")

    # 计算正确率
    if len(predicts) != 0:
        correct_rate = (common_predict - 1) / (len(predicts)+1) if common_predict > 1 else 0  # 避免负值
        print(f"predict寻错的正确率：{correct_rate}")

    return intersection_rate, correct_rate


if __name__ == '__main__':
    whitelist_filename = r"D:\python\python-Project1\项目\test-trainer\Utils\while_list.txt"
    read_filename = r"D:\python\python-Project1\项目\test-trainer\Utils\111.docx"
    openbert1(read_filename, whitelist_filename)

